from docx import Document
import tomllib

res = tomllib.loads("".join(open("./content/_index.md").readlines()[1:-2]))
document = Document()

document.add_heading("{} {} {}".format(res.get("first_name"), res.get("middle_name"), res.get("last_name")), 0)

contact = document.add_table(rows=2, cols=6)
contact_cells = contact.rows[0].cells
contact_cells[0].text = "linkedin"
contact_cells[1].text = "Github"
contact_cells[2].text = "Keybase"
contact_cells[3].text = "Telephone"
contact_cells[4].text = "E-mail"
contact_cells[5].text = "WWW"

contact_cells = contact.rows[1].cells
contact_cells[0].text = res.get("linkedin")
contact_cells[1].text = res.get("github")
contact_cells[2].text = res.get("keybase")
contact_cells[3].text = res.get("telephone")
contact_cells[4].text = res.get("email")
contact_cells[5].text = res.get("web")

quote = document.add_paragraph(res.get("quote"))
quote.italic = True

for index in res.get("page"):
    for index_2 in index.keys():
        if index_2 == "expertise":
            exp = index.get("expertise")
            document.add_heading('Expertise', level=1)
            expertise = document.add_table(rows=len(exp), cols=2)
            for index_3, nr in zip(exp, range(len(exp))):
                expertise.rows[nr].cells[0].add_paragraph(index_3.get("category"))
                expertise.rows[nr].cells[1].add_paragraph(index_3.get("skills"))
            break

document.add_heading('History', level=1)
for index in res.get("page"):
    for index_2 in index.keys():
        if index_2 == "history":
            hist = index.get("history")
            for index_3 in hist:
                document.add_heading(index_3.get("company_name"), level=2)
                heading = document.add_table(rows=2, cols=2)
                heading.rows[0].cells[0].text = index_3.get("job_title")

document.add_heading('Certifications', level=1)
document.add_heading('Presentations', level=1)
document.add_heading('Activities', level=1)

document.save('static/dl/cv.docx')
